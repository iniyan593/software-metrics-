import sys
import subprocess

# Auto-install python-docx if not installed
try:
    import docx
except ImportError:
    print("📦 Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_ewss_word_doc():
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Color Palette Constants
    COLOR_PRIMARY = RGBColor(16, 185, 129)    # Emerald Green (#10B981)
    COLOR_SECONDARY = RGBColor(15, 23, 42)    # Slate Dark (#0F172A)
    COLOR_MUTED = RGBColor(100, 116, 139)     # Slate Muted (#64748B)

    # Helper function for setting table header shading
    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper function for adding styled headings
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        return p

    # --- DOCUMENT CONTENT GENERATION ---

    add_title("🌱 EWSS 2.0 — Sustainability & Effluent Water Assessment Platform")
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Executive Summary, Architecture, Mathematical Scoring Engine & Review Guide")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Executive Summary
    add_heading_1("1. Executive Summary & Project Purpose")
    add_p("The Effluent Water Sustainability & Reuse Score (EWSS 2.0) platform is an enterprise-grade decision-support system designed for ethanol distilleries, industrial plants, and environmental regulatory bodies (such as CPCB / SPCB).")
    add_p("Rather than measuring user satisfaction, EWSS 2.0 measures regulatory compliance, environmental safety, and groundwater depletion risk. It fuses real-time distillery effluent telemetry with external environmental feeds from the Central Water Commission (CWC) and the Central Ground Water Board (CGWB), applying multi-criteria decision modeling (AHP-TOPSIS) and privacy-preserving Federated Learning (FedAvg).")

    # 2. Technical Architecture Table
    add_heading_1("2. Core Technological Components Required for Development")
    add_p("The platform architecture is structured across five core operational layers:")

    t1_data = [
        ["Layer", "Technologies Used", "Core Responsibilities"],
        ["1. Frontend & UI", "Streamlit, Plotly Express", "Interactive dashboard, digital twin sliders, batch trend charts, risk badges."],
        ["2. Data Fusion", "Pandas, NumPy", "Ingestion of company CSV files, state node telemetry fusion, data cleaning."],
        ["3. Scoring Engine", "AHP, TOPSIS Normalization", "Parameter weighting & standardizing physical units into 0.0 - 1.0 sub-scores."],
        ["4. Explainability", "SHAP Attribution, Rules", "Quantifying exact point deductions per parameter & automated alerts."],
        ["5. Privacy & Export", "FedAvg, ReportLab", "Multi-state model aggregation without raw data sharing; PDF & CSV exports."]
    ]

    t1 = doc.add_table(rows=len(t1_data), cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(t1_data):
        for c_idx, cell_value in enumerate(row):
            cell = t1.cell(r_idx, c_idx)
            cell.text = cell_value
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx == 0:
                set_cell_background(cell, "0F172A")
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")

    # 3. Parameters
    add_heading_1("3. Required Input Parameters")
    add_p("The system evaluates 7 critical parameters categorized into Effluent Chemistry, Operational Efficiency, and Environmental Hydrology:")
    
    add_heading_2("A. Effluent Chemistry Parameters (Distillery Telemetry)")
    add_bullet("Range: 4.0 – 10.0 (Optimal: 6.5 – 8.5). Extreme acidity or alkalinity damages aquatic ecosystems.", "1. pH Level: ")
    add_bullet("Range: 10,000 – 120,000 mg/L (Optimal: <= 15,000 mg/L). High chemical oxygen demand.", "2. COD (Chemical Oxygen Demand): ")
    add_bullet("Range: 5,000 – 65,000 mg/L (Optimal: <= 6,000 mg/L). High organic load depletes dissolved oxygen.", "3. BOD (Biochemical Oxygen Demand): ")
    add_bullet("Range: 500 – 6,000 mg/L (Optimal: <= 1,000 mg/L). Elevated salts prevent agricultural effluent reuse.", "4. TDS (Total Dissolved Solids): ")

    add_heading_2("B. Plant Operational Efficiency Parameter")
    add_bullet("Range: 4.0 – 20.0 L water / L ethanol (Optimal: <= 6.5 L/L). Measures water intensity.", "5. Water Consumption Ratio (WCR): ")

    add_heading_2("C. Environmental & Hydrological Feeds (External CWC / CGWB Feeds)")
    add_bullet("Range: 2.0 – 35.0 m bgl (Optimal: <= 8.0 m). Deeper water tables indicate localized aquifer stress.", "6. Groundwater Depth (GWD): ")
    add_bullet("Range: 0.0 – 100.0 mm (Optimal: >= 30.0 mm). Provides natural aquifer recharge and dilution.", "7. Daily Rainfall (RF): ")

    # 4. Mathematical Methodology
    add_heading_1("4. Mathematical Methodology & EWSS Calculation")
    add_p("The composite score is calculated through a 4-step quantitative pipeline:")

    add_heading_2("Step 1: TOPSIS-Inspired Sub-Score Normalization (S_i in [0.0, 1.0])")
    add_p("• For pH (Midpoint target = 7.0): S_pH = max(0, 1 - |pH - 7.0| / 3.0)")
    add_p("• For Cost Criteria (COD, BOD, TDS, Water Ratio, Groundwater Depth): S_i = max(0.0, min(1.0, (Max - x_i) / (Max - Min)))")
    add_p("• For Benefit Criteria (Rainfall): S_Rainfall = max(0.0, min(1.0, x_Rainfall / Target))")

    add_heading_2("Step 2: AHP Weight Vector Allocation")
    
    t2_data = [
        ["Parameter", "Symbol", "AHP Weight", "Weight %"],
        ["Chemical Oxygen Demand", "COD", "0.25", "25%"],
        ["Biochemical Oxygen Demand", "BOD", "0.22", "22%"],
        ["Water Consumption Ratio", "WCR", "0.13", "13%"],
        ["Total Dissolved Solids", "TDS", "0.12", "12%"],
        ["Groundwater Depth", "GWD", "0.10", "10%"],
        ["Daily Rainfall", "RF", "0.10", "10%"],
        ["pH Level", "pH", "0.08", "8%"],
        ["Total Composite Sum", "Sum", "1.00", "100%"]
    ]

    t2 = doc.add_table(rows=len(t2_data), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(t2_data):
        for c_idx, cell_value in enumerate(row):
            cell = t2.cell(r_idx, c_idx)
            cell.text = cell_value
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if r_idx == 0:
                set_cell_background(cell, "10B981")
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            elif r_idx == len(t2_data) - 1:
                set_cell_background(cell, "E2E8F0")
                p.runs[0].font.bold = True

    add_heading_2("Step 3: Composite Score & 95% Confidence Bounds Calculation")
    add_p("• EWSS Composite Score = 100 * SUM( w_i * S_i )")
    add_p("• Combined Standard Deviation = SQRT( SUM( (w_i * sigma_i)^2 ) )")
    add_p("• 95% Margin of Error (MoE) = 1.96 * Combined Standard Deviation * 100")
    add_p("• Confidence Interval = [ EWSS - MoE,  EWSS + MoE ]")

    # 5. Status Classifications
    add_heading_1("5. System Outputs & Status Classifications")
    
    t3_data = [
        ["Score Range", "Status Tier", "Interpretation", "Action Required"],
        [">= 80.0", "EXCELLENT", "Full compliance; safe effluent reuse; zero aquifer stress.", "Maintain standard operations."],
        ["60.0 - 79.9", "ACCEPTABLE", "Moderate operational risk; parameters approaching limits.", "Perform preventive maintenance."],
        ["< 60.0", "CRITICAL RISK", "Non-compliant; severe aquifer stress & environmental pollution.", "Mandatory operational intervention."]
    ]

    t3 = doc.add_table(rows=len(t3_data), cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(t3_data):
        for c_idx, cell_value in enumerate(row):
            cell = t3.cell(r_idx, c_idx)
            cell.text = cell_value
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx == 0:
                set_cell_background(cell, "0F172A")
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Save output
    output_filename = "EWSS_2.0_Project_Summary_Document.docx"
    doc.save(output_filename)
    print(f"✅ Created Word Document successfully: {output_filename}")

if __name__ == "__main__":
    create_ewss_word_doc()